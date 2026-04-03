"""
文件文本提取服务 - 从 Supabase Storage 下载文件并提取纯文本
支持：txt、docx（python-docx），PDF/图片暂时返回提示（OCR 后续接入）
"""
import logging
import io
from typing import Optional

import httpx

logger = logging.getLogger(__name__)


async def download_file_bytes(signed_url: str) -> bytes:
    """下载 Supabase Storage 中的文件内容"""
    async with httpx.AsyncClient(timeout=300.0) as client:
        resp = await client.get(signed_url)
        resp.raise_for_status()
        return resp.content


async def extract_text_from_file(
    file_bytes: bytes,
    file_name: str,
    mime_type: Optional[str] = None,
) -> str:
    """
    从文件内容中提取纯文本 (异步非阻塞包装)
    """
    import asyncio
    return await asyncio.to_thread(_extract_text_from_file_sync, file_bytes, file_name, mime_type)

def _extract_text_from_file_sync(
    file_bytes: bytes,
    file_name: str,
    mime_type: Optional[str] = None,
) -> str:
    name_lower = file_name.lower()

    if name_lower.endswith(".txt") or (mime_type and "text/plain" in mime_type):
        for enc in ("utf-8", "gbk", "gb2312", "utf-16"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if name_lower.endswith(".docx"):
        try:
            import docx  # python-docx
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            
            full_text_parts = []
            
            # 1. 提取所有段落
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text_parts.append(p.text)
            
            # 2. 提取所有表格内容
            for table in doc.tables:
                for row in table.rows:
                    # 将行内单元格合并，模拟表格布局
                    row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_content:
                        full_text_parts.append(" | ".join(row_content))
            
            return "\n".join(full_text_parts)
        except ImportError:
            logger.warning("python-docx not installed, cannot extract docx text")
            return ""
        except Exception as e:
            logger.error(f"docx extraction error: {e}")
            return ""

    # PDF - 使用 PyMuPDF 直接提取文本
    if name_lower.endswith(".pdf") or (mime_type and "pdf" in mime_type):
        try:
            import fitz  # PyMuPDF
            import io
            doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            doc.close()
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("PyMuPDF not installed, cannot extract PDF text")
            return ""
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""

    if any(name_lower.endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]):
        try:
            import fitz  # PyMuPDF
            import io
            # 尝试使用PyMuPDF从图片中提取文本
            doc = fitz.open(stream=io.BytesIO(file_bytes), filetype=name_lower.split('.')[-1])
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            doc.close()
            text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(text)} characters from image: {file_name}")
            return text
        except ImportError:
            logger.warning("PyMuPDF not installed, cannot extract image text")
            return ""
        except Exception as e:
            logger.error(f"Image text extraction error: {e}")
            return ""

    return ""
